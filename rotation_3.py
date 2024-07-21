import pandas as pd
import matplotlib.pyplot as plt
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches

# Constants
calibration_factor = 0.0154
area = 0.000531  # in m^2

# File path
file_path = 'rotation_time.xlsx'


import pandas as pd

# Load the Excel file
file_path = 'Female_new_GLI.xlsx'


from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score

# Separate targets
y_fev1 = df['fev1_predicted']
y_fvc = df['fvc_predicted']

# Train the model for FEV1
model_fev1 = RandomForestRegressor(n_estimators=100, random_state=42)
model_fev1.fit(X_poly, y_fev1)

# Train the model for FVC
model_fvc = RandomForestRegressor(n_estimators=100, random_state=42)
model_fvc.fit(X_poly, y_fvc)

# Predictions
fev1_pred = model_fev1.predict(X_poly)
fvc_pred = model_fvc.predict(X_poly)


# Read the Excel file into a DataFrame, skipping the first row
df = pd.read_excel(file_path, header=None)

# Extract id and date from the first two rows and first column
test_id = df.iloc[0, 1]
test_date = df.iloc[-2, 0]
temp=df.iloc[-2,1]
humidity=df.iloc[-2,2]
pressure=df.iloc[-2,3]

# Transpose DataFrame to have columns as rows
df = df.T

# Set the first row as header and keep only the data rows
df.columns = df.iloc[5]
df = df[5:]

# Select the last two columns (Exhale and Inhale)
df = df.iloc[:, [-2, -1]]
df.columns = ['Exhale', 'Inhale']

# Process the first 4 cells in each column
for col in df.columns:
    for i in range(3):
        if df.iloc[i][col] < 3:
            df.iloc[i][col] = None
        else:
            break
    df[col] = df[col].dropna().reset_index(drop=True)
    
# Add a row of zeros at the beginning
zero_row = pd.DataFrame({'Exhale': [0], 'Inhale': [0]})
df = pd.concat([zero_row, df]).reset_index(drop=True)

# import pandas as pd

# Set pandas options to display the full DataFrame
pd.set_option('display.max_rows', None)  # No limit on the number of rows displayed
pd.set_option('display.max_columns', None)  # No limit on the number of columns
pd.set_option('display.width', None)  # No restriction on the total display width
pd.set_option('display.max_colwidth', None)  # Display full content of each column

# Now, print the full DataFrame
print(df)

# Generate the time column
df['time'] = [i * 0.1 for i in range(len(df))]

# Convert columns to numeric
df['Exhale'] = pd.to_numeric(df['Exhale'], errors='coerce')
df['Inhale'] = pd.to_numeric(df['Inhale'], errors='coerce')
df['time'] = pd.to_numeric(df['time'], errors='coerce')

# Calculate the flow rate for Exhale and Inhale (in m^3/sec)
df['Flowrate_Exhale'] = ((df['Exhale'] * calibration_factor / 0.1) * area) * 1000
df['Flowrate_Inhale'] = ((df['Inhale'] * calibration_factor / 0.1) * area) * 1000

# Calculate volume (in L)
df['Volume_Exhale'] = df['Flowrate_Exhale'] * 0.1
df['Volume_Inhale'] = df['Flowrate_Inhale'] * 0.1

# Calculate the cumulative volume for Exhale and Inhale
df['Cumulative_Volume_Exhale'] = df['Volume_Exhale'].cumsum()
df['Cumulative_Volume_Inhale'] = -df['Volume_Inhale'].cumsum() + df['Cumulative_Volume_Exhale'].iloc[-1]

# Calculate FEV1, FVC, and PEF
FEV1 = df.loc[df['time'] <= 1, 'Volume_Exhale'].sum()
FVC = df['Volume_Exhale'].sum()
PEF = df['Flowrate_Exhale'].max()

# Calculate FIVC (Forced Inspiratory Vital Capacity)
FIVC = df['Volume_Inhale'].sum()

# Calculate additional spirometry parameters
FEV05 = df.loc[df['time'] <= 0.5, 'Volume_Exhale'].sum()
FEV3 = df.loc[df['time'] <= 3, 'Volume_Exhale'].sum()
FEV6 = df.loc[df['time'] <= 6, 'Volume_Exhale'].sum()

FEF25 = df.loc[df['Cumulative_Volume_Exhale'] >= 0.25 * FVC, 'Flowrate_Exhale'].iloc[0]
FEF50 = df.loc[df['Cumulative_Volume_Exhale'] >= 0.5 * FVC, 'Flowrate_Exhale'].iloc[0]
FEF75 = df.loc[df['Cumulative_Volume_Exhale'] >= 0.75 * FVC, 'Flowrate_Exhale'].iloc[0]

FEV3_FVC = FEV3 / FVC
FEV05_FVC = FEV05 / FVC

# Plot the graphs
plt.figure(figsize=(12, 6))
plt.plot(df['Cumulative_Volume_Exhale'], df['Flowrate_Exhale'], label='Flowrate Exhale', color='blue')
plt.plot(df['Cumulative_Volume_Inhale'], -df['Flowrate_Inhale'], label='Flowrate Inhale', color='green')
plt.xlabel('Volume')
plt.ylabel('Flowrate')
plt.title('Flowrate and Cumulative Volume over Time')
plt.legend()
plt.grid(True)
plt.savefig('graph.png')  # Save the graph as a PNG file
plt.show()

plt.figure(figsize=(12, 6))
plt.plot(df['time'], df['Cumulative_Volume_Exhale'], label='Cumulative Volume Exhale', color='blue')
plt.xlabel('Time')
plt.ylabel('Volume')
plt.title('Cumulative Volume over Time')
plt.legend()
plt.grid(True)
plt.savefig('graph2.png')  # Save the graph as a PNG file
plt.show()

# Function to insert the graph into a Word document
def insert_graph_into_word():
    doc = Document()
    doc.add_picture('15.png', width=Inches(2.5))
    doc.add_paragraph(f'ID: {test_id}\nDate: {test_date}', style='BodyText')
    doc.add_heading('Spirometry Results', level=1)
    doc.add_paragraph('The following graph visualizes the flowrate and cumulative volume over time.')
    doc.add_picture('graph.png', width=Inches(5.0))
    doc.add_picture('graph2.png', width=Inches(5.0))
    
    doc.add_paragraph('The following table lists the calculated spirometry parameters.')

    # Add a table to the document
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Add the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Value'

    # Add the spirometry parameters to the table
    parameters = [
        ('FEV1 (L)', f"{FEV1:.2f}"),
        ('FVC (L)', f"{FVC:.2f}"),
        ('PEF (L/sec)', f"{PEF:.2f}"),
        ('FIVC (L)', f"{FIVC:.2f}"),
        ('FEV1/FVC', f"{FEV1/FVC:.2f}"),
        ('FEV.5 (L)', f"{FEV05:.2f}"),
        ('FEV3 (L)', f"{FEV3:.2f}"),
        ('FEV6 (L)', f"{FEV6:.2f}"),
        ('FEF 25% (L/sec)', f"{FEF25:.2f}"),
        ('FEF 50% (L/sec)', f"{FEF50:.2f}"),
        ('FEF 75% (L/sec)', f"{FEF75:.2f}"),
        ('FEV3/FVC', f"{FEV3_FVC:.2f}"),
        ('FEV.5/FVC', f"{FEV05_FVC:.2f}")
    ]

    for param, value in parameters:
        row_cells = table.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = value

    doc.add_paragraph('\n')

    # Save the document
    doc.save('GraphDocument.docx')

# Run the function to create the graph and save it in a Word document
insert_graph_into_word()
