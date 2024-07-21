import pandas as pd
import matplotlib.pyplot as plt
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches
import math
from sklearn.preprocessing import PolynomialFeatures
from sklearn.ensemble import RandomForestRegressor
import time

# Load the Excel file female
file_for_pred_female = 'Female_new_GLI.xlsx'
df_pred_female = pd.read_excel(file_for_pred_female)

column_fev1_f = df_pred_female.iloc[1:, 4]
std_dev_fev1_f = column_fev1_f.std()

column_fvc_f = df_pred_female.iloc[1:, 5]
std_dev_fvc_f = column_fvc_f.std()

column_fev1_fvc_f = df_pred_female.iloc[1:, 6]
std_dev_fev1_fvc_f = column_fev1_fvc_f.std()


# Load the Excel file male
file_for_pred_male = 'Male_new_GLI.xlsx'
df_pred_male = pd.read_excel(file_for_pred_male)

column_fev1_m = df_pred_male.iloc[1:, 4]
std_dev_fev1_m = column_fev1_m.std()

column_fvc_m = df_pred_female.iloc[1:, 5]
std_dev_fvc_m = column_fvc_m.std()

column_fev1_fvc_m = df_pred_female.iloc[1:, 6]
std_dev_fev1_fvc_m = column_fev1_fvc_m.std()



# Check if 'fev1_predicted' and 'fvc_predicted' are present
if 'fev1_predicted' not in df_pred_female.columns or 'fvc_predicted' not in df_pred_female.columns:
    print("Required columns 'fev1_predicted' or 'fvc_predicted' are missing.")
else:
    # Generate polynomial features (degree can be tuned)
    poly = PolynomialFeatures(degree=2, include_bias=False)
    X_poly_female = poly.fit_transform(df_pred_female[['age', 'height']])

    # Separate targets
    y_fev1_female = df_pred_female['fev1_predicted']
    y_fvc_female = df_pred_female['fvc_predicted']

    # Train the model for FEV1
    model_fev1_female = RandomForestRegressor(n_estimators=100, random_state=42)
    model_fev1_female.fit(X_poly_female, y_fev1_female)

    # Train the model for FVC
    model_fvc_female = RandomForestRegressor(n_estimators=100, random_state=42)
    model_fvc_female.fit(X_poly_female, y_fvc_female)

    # Predictions
    fev1_pred_female = model_fev1_female.predict(X_poly_female)
    fvc_pred_female = model_fvc_female.predict(X_poly_female)
    
    
    # Check if 'fev1_predicted' and 'fvc_predicted' are present
if 'fev1_predicted' not in df_pred_male.columns or 'fvc_predicted' not in df_pred_male.columns:
    print("Required columns 'fev1_predicted' or 'fvc_predicted' are missing.")
else:
    # Generate polynomial features (degree can be tuned)
    poly = PolynomialFeatures(degree=2, include_bias=False)
    X_poly_male = poly.fit_transform(df_pred_male[['age', 'height']])

    # Separate targets
    y_fev1_male = df_pred_male['fev1_predicted']
    y_fvc_male = df_pred_male['fvc_predicted']

    # Train the model for FEV1
    model_fev1_male = RandomForestRegressor(n_estimators=100, random_state=42)
    model_fev1_male.fit(X_poly_male, y_fev1_male)

    # Train the model for FVC
    model_fvc_male = RandomForestRegressor(n_estimators=100, random_state=42)
    model_fvc_male.fit(X_poly_male, y_fvc_male)

    # Predictions
    fev1_pred_male = model_fev1_male.predict(X_poly_male)
    fvc_pred_fale = model_fvc_male.predict(X_poly_male)
    

    def predict_fev1_fvc(age, height, gender):
            # Create a DataFrame for the new input
        new_data = pd.DataFrame({'age': [age], 'height': [height]})
        
        if gender == 'F':
            # Generate polynomial features for the new input
            new_data_poly = poly.transform(new_data)
            start_time = time.time()
            # Predict FEV1 and FVC
            fev1_prediction = model_fev1_female.predict(new_data_poly)
            fvc_prediction = model_fvc_female.predict(new_data_poly)    
            
        else:
            # Generate polynomial features for the new input
            new_data_poly = poly.transform(new_data)
            start_time = time.time()
            # Predict FEV1 and FVC
            fev1_prediction = model_fev1_male.predict(new_data_poly)    
            fvc_prediction = model_fvc_male.predict(new_data_poly)  

        end_time = time.time()

        prediction_time = end_time - start_time
        print(f"Time taken for prediction: {prediction_time:.6f} seconds")

        return fev1_prediction[0], fvc_prediction[0]

    
    gender = input("Enter Sex (M/F): ")
    age = int(input("Enter age: "))
    height = int(input("Enter height in cm: "))

    fev1_pred, fvc_pred = predict_fev1_fvc(age, height, gender)
    
    fev1_fvc = fev1_pred/fvc_pred
    
    
    pef_pred = (-1.454 * (age) + 2.368 * (height))/60
    

    # Constants
    area = 0.000531  # in m^2
    # File path
    file_path = 'rotation_time.xlsx'

    # Read the Excel file into a DataFrame, skipping the first row
    df = pd.read_excel(file_path, header=None)

    # Extract id and date from the first two rows and first column
    test_id = df.iloc[0, 1]
    test_date = df.iloc[-2, 0]
    temp = df.iloc[-2, 1]
    humidity = df.iloc[-2, 2]
    pressure = df.iloc[-2, 3]

    # Transpose DataFrame to have columns as rows
    df = df.T

    # Set the first row as header and keep only the data rows
    df.columns = df.iloc[5]
    df = df[5:]

    # Select the last two columns (Exhale and Inhale) and convert to lists
    exhale = df.iloc[:, -2].tolist()
    inhale = df.iloc[:, -1].tolist()

    # Function to remove rows with values <= 3 and limit removal to 4 rows
    def remove_and_shift(column):
        count = 0
        new_column = [0]
        for i, value in enumerate(column):
            if i < 4 and value <= 3 and count < 4:
                count += 1
            else:
                new_column.append(value)

        return new_column

    # Apply the function to each list independently
    exhale_clean = remove_and_shift(exhale)
    inhale_clean = remove_and_shift(inhale)

    print((exhale_clean))
    print((inhale_clean))

    # Generate time lists separately for Exhale and Inhale
    time_exhale = [i * 0.1 for i in range(len(exhale_clean))]
    time_inhale = [i * 0.1 for i in range(len(inhale_clean))]

    # Calculate the flow rate for Exhale and Inhale (in m^3/sec)
    flowrate_exhale = [((-0.0026 * value * value + 0.3187 * value - 0.097) * area) * 1000 for value in exhale_clean]
    flowrate_inhale = [((-0.0026 * value * value + 0.3187 * value - 0.097) * area) * 1000 for value in inhale_clean]

    # Calculate volume (in L)
    volume_exhale = [flow * 0.1 for flow in flowrate_exhale]
    volume_inhale = [flow * 0.1 for flow in flowrate_inhale]

    # Calculate the cumulative volume for Exhale and Inhale
    cumulative_volume_exhale = [sum(volume_exhale[:i+1]) for i in range(len(volume_exhale))]
    cumulative_volume_inhale = [-sum(volume_inhale[:i+1]) + cumulative_volume_exhale[-1] for i in range(len(volume_inhale))]

    # Calculate FEV1, FVC, and PEF
    FEV1 = sum([vol for vol, t in zip(volume_exhale, time_exhale) if t <= 1])
    FVC = sum(volume_exhale)
    PEF = max(flowrate_exhale)

    # Calculate FIVC (Forced Inspiratory Vital Capacity)
    FIVC = sum(volume_inhale)

    # Calculate additional spirometry parameters
    FEV05 = sum([vol for vol, t in zip(volume_exhale, time_exhale) if t <= 0.5])
    FEV3 = sum([vol for vol, t in zip(volume_exhale, time_exhale) if t <= 3])
    FEV6 = sum([vol for vol, t in zip(volume_exhale, time_exhale) if t <= 6])

    # Find flow rates at specific cumulative volumes
    FEF25 = next(flow for flow, vol in zip(flowrate_exhale, cumulative_volume_exhale) if vol >= 0.25 * FVC)
    FEF50 = next(flow for flow, vol in zip(flowrate_exhale, cumulative_volume_exhale) if vol >= 0.5 * FVC)
    FEF75 = next(flow for flow, vol in zip(flowrate_exhale, cumulative_volume_exhale) if vol >= 0.75 * FVC)

    # Calculate FEV1/FVC and FEV0.5/FVC
    FEV3_FVC = FEV3 / FVC
    FEV05_FVC = FEV05 / FVC

    # Calculate mean flow rate between 25% and 75% of FVC
    FEF25_75 = sum([flow for flow, vol in zip(flowrate_exhale, cumulative_volume_exhale) if 0.25 * FVC <= vol <= 0.75 * FVC]) / len([vol for vol in cumulative_volume_exhale if 0.25 * FVC <= vol <= 0.75 * FVC])

    # Calculate peak inspiratory flow rate
    PIFR = max(flowrate_inhale)

    # Calculate mean flow rate between 0.2 and 1.2 liters of FVC
    FEF0_2_1_2 = sum([flow for flow, vol in zip(flowrate_exhale, cumulative_volume_exhale) if 0.2 <= vol <= 1.2]) / len([vol for vol in cumulative_volume_exhale if 0.2 <= vol <= 1.2])

    # Find flow rates at specific cumulative volumes for inspiratory flow
    FIF25 = next(flow for flow, vol in zip(flowrate_inhale[::-1], cumulative_volume_inhale[::-1]) if vol >= 0.25 * FIVC)
    FIF50 = next(flow for flow, vol in zip(flowrate_inhale[::-1], cumulative_volume_inhale[::-1]) if vol >= 0.5 * FIVC)
    FIF75 = next(flow for flow, vol in zip(flowrate_inhale[::-1], cumulative_volume_inhale[::-1]) if vol >= 0.75 * FIVC)


    # Predicted_percentage
    fev1_pred_percentage = (FEV1 / fev1_pred) * 100 # in percentage
    fvc_pred_percentage = (FVC / fvc_pred) * 100 # in percentage  
    pef_pred_percentage = (PEF / pef_pred) * 100 # in percentage    
    fev1_fvc_pred_percentage = (FEV1/FVC / fev1_fvc) * 100 # in percentage
    
    
    if(gender == 'F'):
        z_score_fev1 = (FEV1 - fev1_pred) / std_dev_fev1_f     
        z_score_fvc = (FVC - fvc_pred) / std_dev_fvc_f
        z_score_fev1_fvc = (FEV1/FVC - fev1_fvc) / std_dev_fev1_fvc_f
        
    else:
        z_score_fev1 = (FEV1 - fev1_pred) / std_dev_fev1_m     
        z_score_fvc = (FVC - fvc_pred) / std_dev_fvc_m  
        z_score_fev1_fvc = (FEV1/FVC - fev1_fvc) / std_dev_fev1_fvc_m  
    
    z_score_fev1_fvc = (FEV1/FVC - fev1_fvc) / std_dev_fev1_f    
    
    # Plot the graphs
    plt.figure(figsize=(7, 8))
    plt.plot(cumulative_volume_exhale, flowrate_exhale, label='Flowrate Exhale', color='blue')
    plt.plot(cumulative_volume_inhale, [-flow for flow in flowrate_inhale], label='Flowrate Inhale', color='green')
    plt.xlabel('Volume')
    plt.ylabel('Flowrate')
    plt.title('Flowrate and Cumulative Volume over Time')
    plt.xticks(range(0, int(max(cumulative_volume_exhale)) + 2))
    plt.legend()
    plt.grid(True)
    plt.savefig('graph.png')  # Save the graph as a PNG file
    plt.show()

    plt.figure(figsize=(12, 6))
    plt.plot(time_exhale, cumulative_volume_exhale, label='Cumulative Volume Exhale', color='blue')
    plt.xlabel('Time')
    plt.ylabel('Volume')
    plt.title('Cumulative Volume over Time')
    plt.yticks(range(0, int(max(cumulative_volume_exhale)) + 1, 1))
    plt.legend()
    plt.grid(True)
    plt.savefig('graph2.png')  # Save the graph as a PNG file
    plt.show()

    # Function to insert the graph into a Word document
    def insert_graph_into_word():
        doc = Document()
        doc.add_paragraph(f'ID: {test_id}\nDate: {test_date}', style='BodyText')
        doc.add_heading('Spirometry Results', level=1)
        doc.add_paragraph(f'Temperature: {temp}   Humidity: {humidity}    Pressure:{pressure}', style='BodyText')
        doc.add_paragraph('The following graph visualizes the flowrate and cumulative volume over time.')
        # Add the first image with specified width and height
        doc.add_picture('graph.png', width=Inches(4.0), height=Inches(4.0))
        # Add the second image with specified width and height
        doc.add_picture('graph2.png', width=Inches(6.0), height=Inches(3.0))

        doc.add_paragraph('The following table lists the calculated spirometry parameters.')

        # Add a table to the document
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        # Add the header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parameter'
        hdr_cells[1].text = 'Value'
        hdr_cells[2].text = 'Predicted Value'
        hdr_cells[3].text = 'Pred_Percentage'
        hdr_cells[4].text = 'Z_Score'

        # Add the spirometry parameters to the table
        parameters = [
            ('FEV1 (L)', f"{FEV1:.2f}", f"{fev1_pred:.2f}", f"{fev1_pred_percentage:.2f}", f"{z_score_fev1:.2f}"),
            ('FVC (L)', f"{FVC:.2f}", f"{fvc_pred:.2f}", f"{fvc_pred_percentage:.2f}", f"{z_score_fvc:.2f}"),
            ('PEF (L/sec)', f"{PEF:.2f}", f"{pef_pred:.2f}", f"{pef_pred_percentage:.2f}", f"-"),
            ('FIVC (L)', f"{FIVC:.2f}", f"{FIVC:.2f}", f"-", f"-"),
            ('FEV1/FVC', f"{FEV1/FVC:.2f}", f"{fev1_fvc:.2f}", f"{fev1_fvc_pred_percentage:.2f}", f"{z_score_fev1_fvc:.2f}"),
            ('FEV.5 (L)', f"{FEV05:.2f}", f"{FEV05:.2f}", f"-", f"-"),
            ('FEV3 (L)', f"{FEV3:.2f}", f"{FEV3:.2f}", f"-", f"-"),
            ('FEV6 (L)', f"{FEV6:.2f}", f"{FEV6:.2f}", f"-", f"-"),
            ('FEF 25% (L/sec)', f"{FEF25:.2f}", f"{FEF25:.2f}", f"-", f"-"),
            ('FEF 50% (L/sec)', f"{FEF50:.2f}", f"{FEF50:.2f}", f"-", f"-"),
            ('FEF 75% (L/sec)', f"{FEF75:.2f}", f"{FEF75:.2f}", f"-", f"-"),
            ('FEV.5/FVC', f"{FEV05_FVC:.2f}", f"{FEV05_FVC:.2f}", f"-", f"-"),
            ('FEF 25%-75% (L/sec)', f"{FEF25_75:.2f}", f"{FEF25_75:.2f}", f"-", f"-"),
            ('FEF 0.2-1.2 (L/sec)', f"{FEF0_2_1_2:.2f}", f"{FEF0_2_1_2:.2f}", f"-", f"-"),
            ('FIF 25% (L/sec)', f"{FIF25:.2f}", f"{FIF25:.2f}", f"-", f"-"),
            ('FIF 50% (L/sec)', f"{FIF50:.2f}", f"{FIF50:.2f}", f"-", f"-"),
            ('FIF 75% (L/sec)', f"{FIF75:.2f}", f"{FIF75:.2f}", f"-", f"-"),
            ('PIFR (L/sec)', f"{PIFR:.2f}", f"{PIFR:.2f}", f"-", f"-")
        ]

        for parameter, value, pred_value,pred_percent, z_score  in parameters:
            row_cells = table.add_row().cells
            row_cells[0].text = parameter
            row_cells[1].text = value
            row_cells[2].text = pred_value
            row_cells[3].text = pred_percent
            row_cells[4].text = z_score 

        doc.save('spirometry_results.docx')

    # Insert the graphs and table into the Word document
    insert_graph_into_word()



