import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches

# File path
file_path = 'spirometry_data.xlsx'

# Read the Excel file into a DataFrame, skipping the first row
df = pd.read_excel(file_path)
# Extract id and date from the first two rows and first column
test_id = df.iloc[0, 0]
test_date= df.iloc[-2, 0]
# Transpose DataFrame to have columns as rows
df = df.T
df.columns = df.iloc[2]
df = df[2:]


# Get the last 4 columns (assuming the first column is time)
df = df.iloc[:, -4:]

# Rename columns for clarity
df.columns = ['volume_exhale', 'flowrate_exhale', 'volume_inhale', 'flowrate_inhale']
print(df)
# Create a time column starting from 0 with an increment of 0.1
time = [i * 0.1 for i in range(len(df))]
df['time'] = time

# Calculate additional spirometry parameters
FEV1 = df.loc[df['time'] <= 1, 'volume_exhale'].iloc[-1]
FVC = df['volume_exhale'].iloc[-1]
PEF = df['flowrate_exhale'].max()
FIVC = df['volume_inhale'].iloc[-1]

FEV05 = df.loc[df['time'] <= 0.5, 'volume_exhale'].iloc[-1]
FEV3 = df.loc[df['time'] <= 3, 'volume_exhale'].iloc[-1]
FEV6 = df.loc[df['time'] <= 6, 'volume_exhale'].iloc[-1]

FEF25 = df.loc[df['volume_exhale'] >= 0.25 * FVC, 'flowrate_exhale'].iloc[0]
FEF50 = df.loc[df['volume_exhale'] >= 0.5 * FVC, 'flowrate_exhale'].iloc[0]
FEF75 = df.loc[df['volume_exhale'] >= 0.75 * FVC, 'flowrate_exhale'].iloc[0]

FEV3_FVC = FEV3 / FVC
FEV05_FVC = FEV05 / FVC

# Plot the graph
plt.figure(figsize=(6, 7))
plt.plot(df['volume_exhale'], df['flowrate_exhale'], label='Flowrate Exhale', color='blue')
plt.plot(df['volume_inhale'], -df['flowrate_inhale'], label='Flowrate Inhale', color='green')
plt.xlabel('Volume (Liters)')
plt.ylabel('Flowrate (Liters/second)')
plt.title('Flowrate and Cumulative Volume over Time')
plt.legend()
plt.grid(True)

# Set x-axis ticks and limits
plt.xticks(np.arange(0, max(df['volume_exhale']) + 1, 1))
plt.xlim(0, max(df['volume_exhale']) + 1)

# Set y-axis ticks and limits
# Calculate y-axis limits
max_flowrate = max(abs(df['flowrate_inhale'].max()), abs(df['flowrate_exhale'].max()))
y_limit = int(max_flowrate) + 1 if max_flowrate % 2 == 0 else int(max_flowrate) + 2
y_ticks = np.arange(-y_limit, y_limit , 2)  # Generate even numbers

# Ensure the zero line is included and maintain a 2-unit distance
if 0 not in y_ticks:
    y_ticks = np.concatenate(([0], y_ticks))
plt.yticks(y_ticks)
plt.ylim(-y_limit, y_limit)

plt.savefig('graph.png')  # Save the graph as a PNG file
plt.show()



# Function to insert the graph into a Word document
def insert_graph_into_word():
    doc = Document()
    doc.add_picture('15.png', width=Inches(2.5))
    doc.add_paragraph(f'ID: {test_id}\nDate: {test_date}', style='BodyText')
    doc.add_heading('Spirometry Results', level=1)
    doc.add_paragraph('The following graphs visualize the flowrate and volume over time.')
    doc.add_picture('graph.png', width=Inches(5.0))
    doc.add_paragraph('The following table lists the calculated spirometry parameters.')

    # Add a table to the document
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Add the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parameter'
    hdr_cells[1].text = 'Value'

    # Add the spirometry parameters to the table
    parameters = [
        ('FEV1 (Liters)', f"{FEV1:.2f}"),
        ('FVC (Liters)', f"{FVC:.2f}"),
        ('PEF (Liters/second)', f"{PEF:.2f}"),
        ('FIVC (Liters)', f"{FIVC:.2f}"),
        ('FEV1/FVC', f"{FEV1/FVC:.2f}"),
        ('FEV0.5 (Liters)', f"{FEV05:.2f}"),
        ('FEV3 (Liters)', f"{FEV3:.2f}"),
        ('FEV6 (Liters)', f"{FEV6:.2f}"),
        ('FEF 25% (Liters/second)', f"{FEF25:.2f}"),
        ('FEF 50% (Liters/second)', f"{FEF50:.2f}"),
        ('FEF 75% (Liters/second)', f"{FEF75:.2f}"),
        ('FEV3/FVC', f"{FEV3_FVC:.2f}"),
        ('FEV0.5/FVC', f"{FEV05_FVC:.2f}"),
    ]

    for param, value in parameters:
        row_cells = table.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = value

    doc.add_paragraph('\n')

    # Save the document
    doc.save('Spirometry_Results.docx')

# Run the function to
insert_graph_into_word()